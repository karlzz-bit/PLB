"""
    操作模式	具体含义
'r'	读取 （默认）
'w'	写入（会先截断之前的内容）
'x'	写入，如果文件已经存在会产生异常
'a'	追加，将内容写入到已有文件的末尾
'b'	二进制模式
't'	文本模式（默认）
'+'	更新（既可以读又可以写）
"""
import csv
import random
import xlrd
import xlwt
from docx import Document
from docx.shared import Cm, Pt
from docx.document import Document as Doc
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

# 打开文件进行写操作，如果文件不存在则会创建
with open('example.txt', 'w') as file:
    # 写入内容到文件
    file.write('Hello, World!\n')
    file.write('This is a test file.\n')

# 打开文件进行追加操作
with open('example.txt', 'a') as file:
    # 追加内容到文件末尾
    file.write('Appending a new line.\n')

# 打开文件进行读操作
with open('example.txt', 'r') as file:
    # 读取文件的所有内容
    content = file.read()
    print('File Content:')
    print(content)

# 打开文件进行逐行读取
with open('example.txt', 'r') as file:
    print('Reading file line by line:')
    for line in file:
        print(line.strip())  # 使用strip()去除每行末尾的换行符

# 写入CSV文件
with open('scores.csv', 'w') as file:
    writer = csv.writer(file)
    writer.writerow(['姓名', '语文', '数学', '英语'])
    names = ['关羽', '张飞', '赵云', '马超', '黄忠']
    for name in names:
        scores = [random.randrange(50, 101) for _ in range(3)]
        scores.insert(0, name)
        writer.writerow(scores)
# 读取CSV文件
with open('scores.csv', 'r') as file:
    reader = csv.reader(file, delimiter='|')
    for data_list in reader:
        print(reader.line_num, end='\t')
        for elem in data_list:
            print(elem, end='\t')
        print()
        
#读取excel文件
# # 使用xlrd模块的open_workbook函数打开指定Excel文件并获得Book对象（工作簿）
# wb = xlrd.open_workbook('阿里巴巴2020年股票数据.xls')
# # 通过Book对象的sheet_names方法可以获取所有表单名称
# sheetnames = wb.sheet_names()
# print(sheetnames)
# # 通过指定的表单名称获取Sheet对象（工作表）
# sheet = wb.sheet_by_name(sheetnames[0])
# # 通过Sheet对象的nrows和ncols属性获取表单的行数和列数
# print(sheet.nrows, sheet.ncols)
# for row in range(sheet.nrows):
#     for col in range(sheet.ncols):
#         # 通过Sheet对象的cell方法获取指定Cell对象（单元格）
#         # 通过Cell对象的value属性获取单元格中的值
#         value = sheet.cell(row, col).value
#         # 对除首行外的其他行进行数据格式化处理
#         if row > 0:
#             # 第1列的xldate类型先转成元组再格式化为“年月日”的格式
#             if col == 0:
#                 # xldate_as_tuple函数的第二个参数只有0和1两个取值
#                 # 其中0代表以1900-01-01为基准的日期，1代表以1904-01-01为基准的日期
#                 value = xlrd.xldate_as_tuple(value, 0)
#                 value = f'{value[0]}年{value[1]:>02d}月{value[2]:>02d}日'
#             # 其他列的number类型处理成小数点后保留两位有效数字的浮点数
#             else:
#                 value = f'{value:.2f}'
#         print(value, end='\t')
#     print()
# # 获取最后一个单元格的数据类型
# # 0 - 空值，1 - 字符串，2 - 数字，3 - 日期，4 - 布尔，5 - 错误
# last_cell_type = sheet.cell_type(sheet.nrows - 1, sheet.ncols - 1)
# print(last_cell_type)
# # 获取第一行的值（列表）
# print(sheet.row_values(0))
# # 获取指定行指定列范围的数据（列表）
# # 第一个参数代表行索引，第二个和第三个参数代表列的开始（含）和结束（不含）索引
# print(sheet.row_slice(3, 0, 5))
#写入excel文件
student_names = ['关羽', '张飞', '赵云', '马超', '黄忠']
scores = [[random.randrange(50, 101) for _ in range(3)] for _ in range(5)]
# 创建工作簿对象（Workbook）
wb = xlwt.Workbook()
# 创建工作表对象（Worksheet）
sheet = wb.add_sheet('一年级二班')
# 添加表头数据
titles = ('姓名', '语文', '数学', '英语')
for index, title in enumerate(titles):
    sheet.write(0, index, title)
# 将学生姓名和考试成绩写入单元格
for row in range(len(scores)):
    sheet.write(row + 1, 0, student_names[row])
    for col in range(len(scores[row])):
        sheet.write(row + 1, col + 1, scores[row][col])
# 保存Excel工作簿
wb.save('考试成绩表.xls')

#Word
# 创建代表Word文档的Doc对象
document = Document()  # type: Doc
# 添加大标题
document.add_heading('快快乐乐学Python', 0)
# 添加段落
p = document.add_paragraph('Python是一门非常流行的编程语言，它')
run = p.add_run('简单')
run.bold = True
run.font.size = Pt(18)
p.add_run('而且')
run = p.add_run('优雅')
run.font.size = Pt(18)
run.underline = True
p.add_run('。')

# 添加一级标题
document.add_heading('Heading, level 1', level=1)
# 添加带样式的段落
document.add_paragraph('Intense quote', style='Intense Quote')
# 添加无序列表
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'second item in ordered list', style='List Bullet'
)
# 添加有序列表
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
document.add_paragraph(
    'second item in ordered list', style='List Number'
)

# 添加图片（注意路径和图片必须要存在）
# document.add_picture('resources/guido.jpg', width=Cm(5.2))

# 添加分节符
document.add_section()

records = (
    ('骆昊', '男', '1995-5-5'),
    ('孙美丽', '女', '1992-2-2')
)
# 添加表格
table = document.add_table(rows=1, cols=3)
table.style = 'Dark List'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '姓名'
hdr_cells[1].text = '性别'
hdr_cells[2].text = '出生日期'
# 为表格添加行
for name, sex, birthday in records:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = sex
    row_cells[2].text = birthday

# 添加分页符
document.add_page_break()

# 保存文档
document.save('demo.docx')

#PDF
pdf_canvas = canvas.Canvas('resources/demo.pdf', pagesize=A4)
width, height = A4

# 绘图
image = canvas.ImageReader('resources/guido.jpg')
pdf_canvas.drawImage(image, 20, height - 395, 250, 375)

# 显示当前页
pdf_canvas.showPage()

# 注册字体文件
pdfmetrics.registerFont(TTFont('Font1', 'resources/fonts/Vera.ttf'))
pdfmetrics.registerFont(TTFont('Font2', 'resources/fonts/青呱石头体.ttf'))

# 写字
pdf_canvas.setFont('Font2', 40)
pdf_canvas.setFillColorRGB(0.9, 0.5, 0.3, 1)
pdf_canvas.drawString(width // 2 - 120, height // 2, '你好，世界！')
pdf_canvas.setFont('Font1', 40)
pdf_canvas.setFillColorRGB(0, 1, 0, 0.5)
pdf_canvas.rotate(18)
pdf_canvas.drawString(250, 250, 'hello, world!')

# 保存
pdf_canvas.save()